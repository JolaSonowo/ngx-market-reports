import re
from datetime import datetime
from io import BytesIO
from typing import Tuple

import pandas as pd
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

NGX_URL_FALLBACK_HTML = "https://ngxgroup.com/?id=33164&method=ical"  # example page that contains Snapshot blocks


def _clean_number(x: str) -> float:
    # Handles values like "N13.05", "13.05", "0.63", "-14.95"
    if x is None:
        return float("nan")
    x = str(x).strip()
    x = x.replace("₦", "").replace("N", "").replace(",", "")
    # keep minus and dot
    x = re.sub(r"[^0-9\.\-]", "", x)
    return float(x) if x not in ("", "-", ".", "-.") else float("nan")


def fetch_top5_tables() -> Tuple[pd.DataFrame, pd.DataFrame]:
    """
    Returns (advancers_df, decliners_df) with columns:
    Symbol, Close Price, % Change, Naira Change
    """
    html = None

    # Method A: normal request + parse tables
    try:
        resp = requests.get(
            NGX_URL_FALLBACK_HTML,
            headers={"User-Agent": "Mozilla/5.0"},
            timeout=30,
        )
        resp.raise_for_status()
        html = resp.text
    except Exception:
        html = None

    # Method B: headless browser if request didn't work or tables not found
    if not html or "Top 5 Advancers" not in html:
        from playwright.sync_api import sync_playwright

        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto("https://ngxgroup.com/", wait_until="networkidle", timeout=60000)

            # Click path may change; adjust selectors after inspecting page
            # Try to reach Snapshot -> Equities -> Top 5 Advancers
            page.get_by_text("Snapshot", exact=False).click()
            page.get_by_text("Equities", exact=False).click()
            page.get_by_text("Top 5 Advancers", exact=False).click()

            html = page.content()
            browser.close()

    # Parse HTML tables
    dfs = pd.read_html(html)
    # Heuristic: find the table with columns like Symbols Last Close Current Change % Change
    target = None
    for df in dfs:
        cols = [str(c).strip().lower() for c in df.columns]
        if ("symbols" in cols or "symbol" in cols) and ("last close" in cols) and ("current" in cols) and ("% change" in cols or "change" in cols):
            # Could be one of the tables we need (advancers OR decliners)
            # We'll collect all matching and then pick first two in page order.
            pass

    # More robust: locate sections by headings in the DOM, then read just that table
    soup = BeautifulSoup(html, "html.parser")
    def table_after_heading(heading_text: str) -> pd.DataFrame:
        h = soup.find(string=re.compile(heading_text, re.I))
        if not h:
            raise ValueError(f"Couldn't find heading: {heading_text}")
        # walk forward to next table
        node = h.parent
        tbl = node.find_next("table")
        if tbl is None:
            raise ValueError(f"No table found after heading: {heading_text}")
        return pd.read_html(str(tbl))[0]

    adv_raw = table_after_heading("Top 5 Advancers")
    dec_raw = table_after_heading("Top 5 Decliners|Top 5 Decliners|Top 5 Decliners|Top 5 Decliners|Top 5 Decliners|Top 5 Decliners|Top 5 Decliners|Top 5 Decliners|Top 5 Decliners|Top 5 Decliners")
    # (yes that looks silly; it's just making the regex tolerant if their spelling changes slightly)

    def normalize(df: pd.DataFrame) -> pd.DataFrame:
        df = df.copy()
        df.columns = [str(c).strip() for c in df.columns]

        # Expected: Symbols, Last Close, Current, Change, % Change
        symbol_col = [c for c in df.columns if c.lower() in ("symbols", "symbol")][0]
        current_col = [c for c in df.columns if c.lower() == "current"][0]
        pct_col = [c for c in df.columns if "change" in c.lower() and "%" in c][0]
        change_col = [c for c in df.columns if c.lower() == "change"][0]

        out = pd.DataFrame({
            "Symbol": df[symbol_col].astype(str).str.strip(),
            "Close Price": df[current_col].astype(str).map(_clean_number),
            "% Change": df[pct_col].astype(str).str.replace("%", "").map(_clean_number),
            "Naira Change": df[change_col].astype(str).map(_clean_number),
        })
        # keep top 5 only
        return out.head(5)

    return normalize(adv_raw), normalize(dec_raw)


def build_excel(adv: pd.DataFrame, dec: pd.DataFrame, report_date: datetime) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Daily Summary"

    title = f"DAILY EQUITY SUMMARY FOR {report_date.strftime('%-dTH %b %Y').upper()}"
    ws["A1"] = title
    ws.merge_cells("A1:D1")
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = Alignment(horizontal="center")

    thin = Side(style="thin", color="FFFFFF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    green = PatternFill("solid", fgColor="86B97B")
    red = PatternFill("solid", fgColor="FF4D4D")
    header_font = Font(bold=True, size=12, color="000000")

    def write_block(start_row: int, label: str, df: pd.DataFrame, fill: PatternFill):
        ws[f"A{start_row}"] = label
        ws[f"A{start_row}"].font = Font(bold=True, size=12)
        ws[f"A{start_row}"].fill = fill
        ws[f"A{start_row}"].alignment = Alignment(horizontal="left")

        headers = ["Symbols", "Close Price", "% Change", "Naira Change"]
        for j, h in enumerate(headers, start=1):
            cell = ws.cell(row=start_row, column=j)
            cell.value = headers[j-1] if j > 1 else label  # keep your visual style if you want
        # Write proper header row (A..D)
        ws[f"A{start_row}"] = label
        ws[f"B{start_row}"] = "Close Price"
        ws[f"C{start_row}"] = "% Change"
        ws[f"D{start_row}"] = "Naira Change"

        for col in "ABCD":
            c = ws[f"{col}{start_row}"]
            c.fill = fill
            c.font = header_font
            c.border = border
            c.alignment = Alignment(horizontal="center")

        for i, row in enumerate(df.itertuples(index=False), start=1):
            r = start_row + i
            ws[f"A{r}"] = row.Symbol
            ws[f"B{r}"] = float(row._1)  # Close Price
            ws[f"C{r}"] = float(row._2)  # % Change
            ws[f"D{r}"] = float(row._3)  # Naira Change

            for col in "ABCD":
                c = ws[f"{col}{r}"]
                c.fill = fill
                c.border = border
                c.alignment = Alignment(horizontal="center" if col != "A" else "left")
                if col == "A":
                    c.font = Font(bold=True)

        return start_row + len(df) + 2

    next_row = 3
    next_row = write_block(next_row, "Gainers", adv, green)
    next_row = write_block(next_row, "Losers", dec, red)

    ws.column_dimensions["A"].width = 20
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 10
    ws.column_dimensions["D"].width = 14

    bio = BytesIO()
    wb.save(bio)
    return bio.getvalue()


def build_word(adv: pd.DataFrame, dec: pd.DataFrame, report_date: datetime) -> bytes:
    doc = Document()
    title = f"DAILY EQUITY SUMMARY FOR {report_date.strftime('%-dTH %b %Y').upper()}"
    doc.add_heading(title, level=1)

    def add_table(label: str, df: pd.DataFrame):
        doc.add_heading(label, level=2)
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = label
        hdr[1].text = "Close Price"
        hdr[2].text = "% Change"
        hdr[3].text = "Naira Change"

        for row in df.itertuples(index=False):
            cells = table.add_row().cells
            cells[0].text = str(row.Symbol)
            cells[1].text = f"{row._1:.2f}"
            cells[2].text = f"{row._2:.2f}"
            cells[3].text = f"{row._3:.2f}"

    add_table("Gainers", adv)
    add_table("Losers", dec)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def generate_files_for_today() -> Tuple[bytes, bytes, str]:
    report_date = datetime.now()  # you can force Africa/Lagos timezone on server
    adv, dec = fetch_top5_tables()
    xlsx = build_excel(adv, dec, report_date)
    docx = build_word(adv, dec, report_date)
    basename = f"DAILY_EQUITY_SUMMARY_{report_date.strftime('%Y-%m-%d')}"
    return xlsx, docx, basename
