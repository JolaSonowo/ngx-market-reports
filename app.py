import datetime
import io
import requests
import pandas as pd
from flask import Flask, render_template, send_file
from docx import Document

app = Flask(__name__)

# --- DATA FETCHING LOGIC ---
def get_ngx_api_data(endpoint):
    """Fetches Top/Bottom symbols directly from NGX API"""
    url = f"https://doclib.ngxgroup.com/REST/api/mrkstat/{endpoint}"
    headers = {
        "User-Agent": "Mozilla/5.0",
        "Referer": "https://ngxgroup.com/"
    }
    try:
        response = requests.get(url, headers=headers, timeout=10)
        data = response.json()
        
        output = []
        for item in data[:5]:  # Get Top 5
            last_close = float(item.get('LAST_CLOSE', 0))
            price_change = float(item.get('PERCENTAGE_CHANGE', 0))
            todays_close = float(item.get('TODAYS_CLOSE', 0))
            
            # Use same math as NGX website logic
            pc_val = (price_change / last_close * 100) if last_close != 0 else 0
            
            output.append({
                "symbol": item.get('SYMBOL', 'N/A'),
                "close": f"N{todays_close:.2f}",
                "n_chg": f"{price_change:.2f}",
                "pc_chg": f"{pc_val:.2f}%",
                "raw_price": todays_close,
                "raw_pct": pc_val
            })
        return output
    except Exception as e:
        print(f"API Error for {endpoint}: {e}")
        return []

# --- ROUTES ---

@app.route('/')
def home():
    gainers = get_ngx_api_data("topsymbols")
    losers = get_ngx_api_data("bottomsymbols")
    today = datetime.date.today().strftime("%dTH %b %Y").upper()
    return render_template('index.html', gainers=gainers, losers=losers, date=today)

@app.route('/download/excel')
def download_excel():
    gainers = get_ngx_api_data("topsymbols")
    losers = get_ngx_api_data("bottomsymbols")
    
    # Create in-memory buffer
    output = io.BytesIO()
    
    # Use Pandas to write to Excel
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        pd.DataFrame(gainers).to_excel(writer, sheet_name='Top Gainers', index=False)
        pd.DataFrame(losers).to_excel(writer, sheet_name='Top Losers', index=False)
    
    output.seek(0)
    return send_file(
        output,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=f"NGX_Report_{datetime.date.today()}.xlsx"
    )

@app.route('/download/word')
def download_word():
    gainers = get_ngx_api_data("topsymbols")
    losers = get_ngx_api_data("bottomsymbols")
    
    doc = Document()
    doc.add_heading(f'NGX Market Report - {datetime.date.today()}', 0)
    
    for title, data in [("Top Gainers", gainers), ("Top Losers", losers)]:
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Symbol'
        hdr_cells[1].text = 'Price'
        hdr_cells[2].text = '% Change'
        
        for item in data:
            row_cells = table.add_row().cells
            row_cells[0].text = item['symbol']
            row_cells[1].text = item['close']
            row_cells[2].text = item['pc_chg']
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return send_file(
        output,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"NGX_Report_{datetime.date.today()}.docx"
    )

if __name__ == '__main__':
    # host='0.0.0.0' allows colleagues on the same Wi-Fi to access via your IP
    app.run(host='0.0.0.0', port=8080, debug=True)
