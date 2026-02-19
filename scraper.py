import os
from datetime import datetime, time as dt_time
from zoneinfo import ZoneInfo
import pandas as pd
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document

REPORT_DIR = "reports"
os.makedirs(REPORT_DIR, exist_ok=True)
NGX_URL = "https://ngxgroup.com/exchange/data/equities-price-list/"
TIMEZONE = ZoneInfo("Africa/Lagos")


def parse_number(value, is_percent=False):
    try:
        val = value.replace(",", "").replace("N", "").replace("--", "0").strip()
        if is_percent:
            val = val.replace("%", "")
        return float(val)
    except:
        return 0.0


def scrape_equities_table():
    options = Options()
    options.binary_location = "/usr/bin/chromium"  # For Render
    options.add_argument("--headless")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    driver = webdriver.Chrome(options=options)
    driver.get(NGX_URL)

    try:
        WebDriverWait(driver, 30).until(
            lambda d: len(d.find_elements(By.CSS_SELECTOR, "table tbody tr")) > 10
        )

        rows = driver.find_elements(By.CSS_SELECTOR, "table tbody tr")
        data = []

        for row in rows:
            cells = row.find_elements(By.TAG_NAME, "td")
            if len(cells) < 4:
                continue
            symbol = cells[0].text.strip()
            current = parse_number(cells[1].text)
            change = parse_number(cells[2].text)
            pct_change = parse_number(cells[3].text, is_percent=True)
            data.append([symbol, current, change, pct_change])

        df = pd.DataFrame(data, columns=["Symbols", "Current", "Change", "% Change"])
        return df

    finally:
        driver.quit()


def get_top_movers(df):
    gainers = df.sort_values("% Change", ascending=False).head(5)
    decliners = df.sort_values("% Change", ascending=True).head(5)
    return gainers, decliners


def generate_excel(gainers, decliners):
    path = os.path.join(REPORT_DIR, "NGX_Market_Report.xlsx")
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        gainers.to_excel(writer, sheet_name="Top Gainers", index=False)
        decliners.to_excel(writer, sheet_name="Top Decliners", index=False)
    return path


def generate_word(gainers, decliners):
    path = os.path.join(REPORT_DIR, "NGX_Market_Report.docx")
    doc = Document()
    doc.add_heading("Nigerian Exchange Market Summary", level=1)
    doc.add_paragraph(f"Generated: {datetime.now(TIMEZONE)}")
    doc.add_heading("Top 5 Gainers", level=2)
    add_table(doc, gainers)
    doc.add_heading("Top 5 Decliners", level=2)
    add_table(doc, decliners)
    doc.save(path)
    return path


def add_table(doc, df):
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = col
    for i, row in df.iterrows():
        for j, val in enumerate(row):
            table.cell(i + 1, j).text = str(val)
